"""Word (.docx) hydraulic report generator — pure Python.

Charts are rendered with matplotlib (no headless browser needed). One public
entry point, ``build_report``, produces either a single-line datasheet or a
full multi-line project report (cover + summary line list + per-line datasheets).
"""
from __future__ import annotations

from io import BytesIO
from datetime import datetime

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from .results import COLUMNS

_BLUE = "2563EB"
_HDR_FG = RGBColor(0xFF, 0xFF, 0xFF)
_ALT = "F1F5F9"
_GREY = RGBColor(0x64, 0x74, 0x8B)


# ── low-level docx helpers ───────────────────────────────────────────────────
def _shade(cell, hex_fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_fill)
    tcPr.append(shd)


def _hdr_cell(cell, text, size=9):
    _shade(cell, _BLUE)
    cell.text = ""
    run = cell.paragraphs[0].add_run(str(text))
    run.font.bold = True
    run.font.color.rgb = _HDR_FG
    run.font.size = Pt(size)


def _cell(cell, text, size=9, bold=False, alt=False):
    if alt:
        _shade(cell, _ALT)
    cell.text = ""
    run = cell.paragraphs[0].add_run("" if text is None else str(text))
    run.font.size = Pt(size)
    run.font.bold = bold


def _kv_table(doc, rows, widths=(2.4, 3.6)):
    t = doc.add_table(rows=len(rows), cols=2)
    t.style = "Table Grid"
    for i, (k, v) in enumerate(rows):
        _cell(t.rows[i].cells[0], k, bold=True, alt=(i % 2 == 1))
        _cell(t.rows[i].cells[1], v, alt=(i % 2 == 1))
    for r in t.rows:
        r.cells[0].width = Inches(widths[0])
        r.cells[1].width = Inches(widths[1])
    return t


# ── matplotlib profile chart ─────────────────────────────────────────────────
def _profile_png(result, width_in=6.3, height_in=3.1) -> BytesIO:
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    x, z, P = result.sketch_x, result.sketch_z, result.sketch_P
    fig, (ax1, ax2) = plt.subplots(2, 1, figsize=(width_in, height_in), sharex=True)
    ax1.plot(x, P, marker="o", color="#2563EB", lw=2)
    ax1.set_ylabel("Pressure (bara)")
    ax1.grid(alpha=0.3)
    ax2.plot(x, z, marker="o", color="#64748b", lw=2)
    ax2.set_ylabel("Elevation (m)")
    ax2.set_xlabel("Cumulative distance (m)")
    ax2.grid(alpha=0.3)
    fig.tight_layout()
    buf = BytesIO()
    fig.savefig(buf, format="png", dpi=150)
    plt.close(fig)
    buf.seek(0)
    return buf


# ── per-line summary stats (for the project line-list table) ─────────────────
def _line_stats(line, result):
    rows = result.rows
    r0 = rows[0] if rows else None
    pipes = [r for r in rows if r.type == "Pipe"]
    dns = {r.pipe for r in pipes if r.pipe}
    max_v = max((r.v_ms for r in pipes), default=0.0)
    max_ratio = max((r.v_over_ve for r in pipes), default=0.0)
    status = "⚠ Review" if result.warnings else "OK"
    return {
        "tag": line.get("tag", ""),
        "service": line.get("service", ""),
        "fluid": r0.fluid if r0 else "",
        "flow_kgh": r0.flow_kgh if r0 else 0.0,
        "dn": next(iter(dns)) if len(dns) == 1 else ("various" if dns else "—"),
        "inlet": result.inlet_P_bara,
        "outlet": result.outlet_P_bara,
        "dp": result.total_dp_kpa,
        "max_v": max_v,
        "max_ratio": max_ratio,
        "status": status,
    }


# ── document building blocks ─────────────────────────────────────────────────
def _cover(doc, meta, n_lines):
    h = doc.add_heading("Line Hydraulic Calculation", level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph(meta.get("project_name") or "Untitled Project")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if sub.runs:
        sub.runs[0].font.size = Pt(14)
        sub.runs[0].font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)
    doc.add_paragraph()
    _kv_table(doc, [
        ("Project", meta.get("project_name", "")),
        ("Client", meta.get("client", "")),
        ("Document", "Line Hydraulic Calculation"),
        ("Lines covered", str(n_lines)),
        ("Revision", meta.get("rev", "")),
        ("Date", meta.get("date", datetime.now().strftime("%Y-%m-%d"))),
        ("Calculated by", meta.get("calc_by", "")),
        ("Checked by", meta.get("checked_by", "")),
    ])
    doc.add_paragraph()


def _methodology(doc, solver_meta, notes):
    doc.add_heading("Method & Assumptions", level=1)
    for b in [
        f"Pressure is marched section by section; fluid properties (density, "
        f"velocity, void fraction) are recomputed at each section's local "
        f"pressure. Two-phase ΔP correlation: {solver_meta.get('correlation','Beggs-Brill')}; "
        f"voidage: {solver_meta.get('voidage','Homogeneous')}.",
        "These are steady incompressible-flow correlations, reliable below ~Mach "
        "0.3 in the gas phase; sections beyond that are flagged. Absolute accuracy "
        "is typically ±20–30 % for two-phase mixtures.",
        "Elevation: ΔP_grav = ρ_insitu · g · Δz uses the geometric elevation "
        "change, separate from friction (which uses the fitting-inclusive "
        "effective length).",
        "Erosion: mixture velocity is checked against the API RP 14E limit (C = 100); "
        "V/V_e > 1 is flagged.",
        "Fitting losses use the equivalent-length method (ΣL_e/D from a standard "
        "fittings table).",
    ]:
        p = doc.add_paragraph(b, style="List Bullet")
        if p.runs:
            p.runs[0].font.size = Pt(9)
    if notes:
        doc.add_paragraph()
        p = doc.add_paragraph("Project-specific notes: " + notes)
        if p.runs:
            p.runs[0].font.size = Pt(9)
            p.runs[0].font.italic = True


def _summary_table(doc, stats_list):
    doc.add_heading("Line List — Summary", level=1)
    cols = ["Tag", "Service", "Fluid", "Flow (kg/h)", "Inlet (bara)",
            "Outlet (bara)", "ΔP (kPa)", "Max V (m/s)", "Max V/V_e", "Status"]
    t = doc.add_table(rows=len(stats_list) + 1, cols=len(cols))
    t.style = "Table Grid"
    for j, c in enumerate(cols):
        _hdr_cell(t.rows[0].cells[j], c, size=8)
    for i, s in enumerate(stats_list, start=1):
        vals = [s["tag"], s["service"], s["fluid"], f"{s['flow_kgh']:.1f}",
                f"{s['inlet']:.3f}", f"{s['outlet']:.3f}", f"{s['dp']:.2f}",
                f"{s['max_v']:.2f}", f"{s['max_ratio']:.3f}", s["status"]]
        for j, v in enumerate(vals):
            _cell(t.rows[i].cells[j], v, size=8, alt=(i % 2 == 0))
    doc.add_paragraph()


def _line_datasheet(doc, line, result, solver_meta):
    tag = line.get("tag") or line["id"]
    svc = line.get("service") or ""
    doc.add_heading(f"Line {tag}" + (f" — {svc}" if svc else ""), level=1)

    r0 = result.rows[0] if result.rows else None
    _kv_table(doc, [
        ("Inlet pressure", f"{result.inlet_P_bara:.4f} bara"),
        ("Outlet pressure", f"{result.outlet_P_bara:.4f} bara"),
        ("Total ΔP", f"{result.total_dp_kpa:.3f} kPa  ({result.total_dp_kpa/100:.4f} bar)"),
        ("  – friction (pipe + fittings)", f"{result.total_dp_fric_kpa:.3f} kPa"),
        ("     · of which fittings/minor losses", f"{sum(r.dp_fit_kpa for r in result.rows):.3f} kPa"),
        ("  – elevation/equipment", f"{result.total_dp_grav_kpa:.3f} kPa"),
        ("Fluid", r0.fluid if r0 else ""),
        ("Composition", r0.composition if r0 else ""),
        ("Flow", f"{r0.flow_kgh:.1f} kg/h" if r0 else ""),
        ("Correlation", solver_meta.get("correlation", "Beggs-Brill")),
    ])
    doc.add_paragraph()

    # Section line-list table (transposed: properties as rows, sections as cols)
    records = [r.to_dict() for r in result.rows]
    props = [c for c in COLUMNS if c != "Element"]
    t = doc.add_table(rows=len(props) + 1, cols=len(records) + 1)
    t.style = "Table Grid"
    _hdr_cell(t.rows[0].cells[0], "Segment →", size=8)
    for j, rec in enumerate(records, start=1):
        _hdr_cell(t.rows[0].cells[j], f"#{j}  {rec['Element']}", size=8)
    for i, prop in enumerate(props, start=1):
        _cell(t.rows[i].cells[0], prop, size=8, bold=True, alt=(i % 2 == 1))
        for j, rec in enumerate(records, start=1):
            _cell(t.rows[i].cells[j], rec[prop], size=8, alt=(i % 2 == 1))
    doc.add_paragraph()

    # Profile chart
    try:
        doc.add_picture(_profile_png(result), width=Inches(6.3))
        cap = doc.add_paragraph("Pressure march and elevation profile.")
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if cap.runs:
            cap.runs[0].font.size = Pt(8)
            cap.runs[0].font.color.rgb = _GREY
    except Exception:
        pass

    if result.warnings:
        p = doc.add_paragraph()
        run = p.add_run("Warnings")
        run.font.bold = True
        run.font.color.rgb = RGBColor(0xB4, 0x53, 0x09)
        for w in result.warnings:
            wp = doc.add_paragraph(w, style="List Bullet")
            if wp.runs:
                wp.runs[0].font.size = Pt(9)
                wp.runs[0].font.color.rgb = RGBColor(0xB4, 0x53, 0x09)


def _footer(doc, meta):
    sec = doc.sections[0]
    p = sec.footer.paragraphs[0]
    p.text = (f"{meta.get('project_name','')}  ·  Rev {meta.get('rev','')}  ·  "
              f"{meta.get('date','')}  ·  Quickpipe")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if p.runs:
        p.runs[0].font.size = Pt(8)
        p.runs[0].font.color.rgb = _GREY


# ── public API ───────────────────────────────────────────────────────────────
def build_report(meta, lines_results, solver_meta, *, single=False) -> BytesIO:
    """lines_results: list of (line_dict, MarchResult). single=True for a
    one-line datasheet (no cover/summary)."""
    doc = Document()
    sec = doc.sections[0]
    sec.left_margin = sec.right_margin = Inches(0.8)
    sec.top_margin = sec.bottom_margin = Inches(0.8)

    if single:
        line, result = lines_results[0]
        h = doc.add_heading("Line Hydraulic Datasheet", level=0)
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _kv_table(doc, [
            ("Project", meta.get("project_name", "")),
            ("Line tag", line.get("tag", "")),
            ("Service", line.get("service", "")),
            ("Rev / Date", f"{meta.get('rev','')}  ·  {meta.get('date','')}"),
            ("Calc by / Checked", f"{meta.get('calc_by','')}  /  {meta.get('checked_by','')}"),
        ])
        doc.add_paragraph()
        _methodology(doc, solver_meta, meta.get("notes", ""))
        _line_datasheet(doc, line, result, solver_meta)
    else:
        _cover(doc, meta, len(lines_results))
        _methodology(doc, solver_meta, meta.get("notes", ""))
        _summary_table(doc, [_line_stats(ln, res) for ln, res in lines_results])
        for ln, res in lines_results:
            doc.add_page_break()
            _line_datasheet(doc, ln, res, solver_meta)

    _footer(doc, meta)
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
